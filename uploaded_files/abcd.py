import sys
sys.stdout.reconfigure(encoding='utf-8')
import os
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from sentence_transformers import SentenceTransformer, util
import io
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from waitress import serve

# Load environment variables
load_dotenv()
FLASK_PORT = os.getenv("FLASK_PORT", 5001)

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

nlp = SentenceTransformer('all-MiniLM-L6-v2')

INTERROGATIVE_WORDS = ["What", "Who", "Why", "How", "Where", "When", "Which", "Whom", "Do", "Does", "Did", "Has", "Have", "Had", "Is", "Was", "Are", "Were", "Am", "Will", "Would", "Can", "Could", "Shall", "Should", "May", "Might"]
QUESTION_KEYWORDS = ["Explain", "Define", "Describe", "Justify", "Discuss", "State the", "Analyze", "Compare", "Elaborate", "_______", "of the passage", "by the passage", "from the passage", "in the passage", "is the passage", "Write", "Find", "Evaluate"]

def extract_questions_from_docx(file_path):
    """Extract questions from a .docx file without altering their form."""
    document = Document(file_path)
    questions = []
    seen_questions = set()  # Track duplicates
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text or text in seen_questions:
            continue  # Skip empty or duplicate questions
        
        # Check if the text is a valid question
        if text.endswith(("?", ":", ":-", "is", "was", "are", "were", "because", "from", "among", "has", "have", "had", "that", "which", "to", "_______", "as", "means", "mean", "meant", "will", "would", "can", "could", "shall", "should", "may", "might")):
            questions.append(text)
            seen_questions.add(text)  # Keep track of unique questions
        elif text.split()[0] in INTERROGATIVE_WORDS or any(keyword in text for keyword in QUESTION_KEYWORDS):
            questions.append(text)
            seen_questions.add(text)
    
    return questions

def find_similar_questions(all_questions, file_mappings, threshold=0.5):
    """Find similar questions, merge overlapping groups, and remove duplicates."""
    question_to_source = {q: file_mappings[i] for i, q in enumerate(all_questions)}
    unique_questions = list(question_to_source.keys())  # Preserve order
    question_embeddings = nlp.encode(unique_questions, convert_to_tensor=True)

    groups = []  # List of merged question groups

    for i, question in enumerate(unique_questions):
        similarities = util.pytorch_cos_sim(question_embeddings[i], question_embeddings)[0]
        group = {question: question_to_source[question]}  # Store the reference question

        for j, similarity in enumerate(similarities):
            if similarity >= threshold and unique_questions[j] != question:
                group[unique_questions[j]] = question_to_source[unique_questions[j]]

        # Merge with existing groups if there’s overlap
        merged = False
        for existing_group in groups:
            if any(q in existing_group for q in group):
                existing_group.update(group)
                merged = True
                break
        
        if not merged:
            groups.append(group)  # Create a new group

    # Convert merged groups into structured output, but only include groups with more than one question
    results = {}
    for group in groups:
        if len(group) > 1:  # Only keep groups that have more than one question
            reference_question = list(group.keys())[0]
            reference_source = group[reference_question]
            results[f"{reference_question} (From {reference_source})"] = [
                {"question": q, "source": s} for q, s in group.items() if q != reference_question
            ]

    return results

@app.route('/upload', methods=['POST'])
def upload_files():
    """Handles file uploads and processes questions."""
    if 'doc-upload' not in request.files:
        return jsonify({'error': 'No files provided'}), 400

    files = request.files.getlist('doc-upload')
    if len(files) < 2:
        return jsonify({'error': 'At least two files are required'}), 400

    file_paths = []
    for file in files:
        if file.filename.endswith('.docx'):
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
            file.save(file_path)
            file_paths.append(file_path)

    if not file_paths:
        return jsonify({'error': 'No valid .docx files uploaded'}), 400

    all_questions = []
    file_mappings = []
    
    for i, file_path in enumerate(file_paths):
        questions = extract_questions_from_docx(file_path)
        all_questions.extend(questions)
        file_mappings.extend([f"Paper {i + 1}"] * len(questions))

    results = find_similar_questions(all_questions, file_mappings, threshold=0.5)
    return jsonify(results)

@app.route('/download', methods=['POST'])
def download_docx():
    """Generates and sends a Word document containing similar questions."""
    data = request.json
    document = Document()

    heading = document.add_heading(level=1)
    heading_run = heading.add_run('Similar Questions Report')
    heading_run.font.color.rgb = RGBColor(0, 0, 0)
    heading_run.font.size = Pt(20)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()

    for ref_question, similar_questions in data.items():
        ref_para = document.add_paragraph()
        ref_run = ref_para.add_run(f"Reference Question - {ref_question}")
        ref_run.bold = True
        ref_run.font.color.rgb = RGBColor(150, 75, 0)
        ref_run.font.size = Pt(13)
        
        similar_heading = document.add_paragraph("Similar Questions:", style='Heading 2')
        for run in similar_heading.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.size = Pt(14)
        
        for i, match in enumerate(similar_questions, start=1):
            document.add_paragraph(
                f"{i}. {match['question']} (From {match['source']})"
            )
        document.add_paragraph()

    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name='similar_questions.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == "__main__":
    serve(app, host="127.0.0.1", port=int(FLASK_PORT))
